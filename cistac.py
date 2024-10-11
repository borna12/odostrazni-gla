from docx import Document

def clean_document(file_path, output_file):
    # Load the document
    doc = Document(file_path)

    # Set to keep track of unique paragraphs
    unique_paragraphs = set()

    # List to store cleaned paragraphs
    cleaned_paragraphs = []

    # Loop through all paragraphs in the document
    for para in doc.paragraphs:
        text = para.text.strip()

        # Check if the paragraph starts with "-" or is a duplicate
        if text.startswith("-") or text in unique_paragraphs:
            continue  # Skip paragraphs that match these conditions

        # If it's not a duplicate or unwanted, add it to unique_paragraphs
        unique_paragraphs.add(text)
        cleaned_paragraphs.append(text)

    # Create a new document to save the cleaned paragraphs
    new_doc = Document()

    # Add the cleaned paragraphs to the new document
    for paragraph in cleaned_paragraphs:
        new_doc.add_paragraph(paragraph)

    # Save the modified document
    new_doc.save(output_file)
    print(f"Document saved as {output_file}")

# File path for the input and output files
input_file = "geslar.docx"  # Replace with your input file
output_file = "nov.docx"

# Call the function
clean_document(input_file, output_file)
